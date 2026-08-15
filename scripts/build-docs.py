from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs"

SCREENS = [
    ("Index", "/", "Directory page for opening every screen in the fixture."),
    ("Candidate assessment welcome", "/welcome", "Assessment start screen."),
    ("Candidate assessment question list", "/questions", "Assessment task list, written response, MCQs, and progress rail."),
    ("Candidate challenge overview", "/challenge", "Single coding challenge overview before the workspace."),
    ("Candidate coding workspace", "/workspace", "Coding editor, output panel, footer, and integrity warning modal."),
    ("Background check dashboard", "/background-check", "Interviewer-side verification dashboard."),
    ("Authentication - Sign in", "/auth/sign-in", "Login form."),
    ("Authentication - Create account", "/auth/create-account", "Registration form."),
    ("Authentication - Forgot password", "/auth/forgot-password", "Password recovery form."),
    ("Authentication - Reset password", "/auth/reset-password", "New password form."),
    ("Authentication - Check your email", "/auth/check-email", "Email confirmation state."),
]

BUGS = {
    "Index": [
        "Not present in the Figma file; candidate must identify it as navigation-only fixture chrome.",
        "Uses card styling and copy that do not belong to the product design system.",
        "Typography and spacing are intentionally more marketing-like than the Figma application UI.",
    ],
    "Candidate assessment welcome": [
        "Header height, padding, logo, and avatar shape differ from Figma.",
        "Title says 'Ready to prove your skills?' instead of the Figma title.",
        "Assessment metadata is wrong: parts, code count, written count, and time limit are changed.",
        "The footer message incorrectly says the assessment can be paused.",
        "Card border, radius, shadow, divider style, and content spacing drift from Figma.",
        "Sidebar rules contain the opposite guidance from the original design.",
    ],
    "Candidate assessment question list": [
        "Coding card has wrong duration, language, action label, and layout proportions.",
        "Written-response prompt and prefilled response do not match Figma.",
        "MCQ choices are intentionally wrong and one incorrect option is preselected.",
        "Progress rail timer, status labels, count, and final action are changed.",
        "Cards use heavier borders, larger radii, and different vertical rhythm than Figma.",
        "The final short question is structurally different from the original MCQ card.",
    ],
    "Candidate challenge overview": [
        "Screen title, eyebrow, challenge name, difficulty, points, and duration are wrong.",
        "Navigation header uses recreated logo text and mismatched spacing.",
        "Brief copy changes the algorithmic requirement.",
        "Topic list differs from Figma's DFS, stack pattern, and input-validation content.",
        "Footer reassurance says autosave is disabled and uses a different button label.",
        "The challenge cards use different widths, radii, and shadow intensity.",
    ],
    "Candidate coding workspace": [
        "Header title is singular and not an exact text match.",
        "Problem panel difficulty, sample inputs, sample outputs, and checklist are wrong.",
        "Editor language, saved state, code body, line count, shortcut, and run button label differ.",
        "Output title, test input, empty state, and result count differ from Figma.",
        "Footer countdown and submit label are changed.",
        "Warning modal has wrong title, warning count, body copy, action order, size, position, and backdrop.",
    ],
    "Background check dashboard": [
        "Header title and sidebar active navigation differ from the Figma dashboard.",
        "Candidate name, role context, status, dates, and upload fields are intentionally changed.",
        "Primary action invites the candidate instead of matching the Figma reminder/upload actions.",
        "Employer table rows include incorrect statuses, outcomes, and dates.",
        "Sidebar profile name and role are wrong.",
        "Layout uses larger pills and rounded controls that do not match the Figma table density.",
    ],
    "Authentication - Sign in": [
        "Left visual panel is recreated with CSS gradients instead of the Figma image asset.",
        "Branding text and tagline are wrong.",
        "Form title, subtitle, labels, placeholders, button copy, and footer copy differ.",
        "Input radius and dimensions do not match the Figma rounded fields.",
        "Panel centering and vertical spacing are intentionally off.",
    ],
    "Authentication - Create account": [
        "Form title, helper text, footer text, and labels differ from Figma.",
        "Form is vertically shifted down compared with the original frame.",
        "Password guidance is intentionally incorrect.",
        "Left panel styling repeats the wrong sign-in treatment rather than matching the Figma asset.",
    ],
    "Authentication - Forgot password": [
        "Uses 'username' instead of the work email field.",
        "Button copy and body copy are not the Figma text.",
        "Form is shifted upward relative to the Figma layout.",
        "The visual panel and spacing reuse the wrong generic auth treatment.",
    ],
    "Authentication - Reset password": [
        "Form is shifted horizontally from the Figma center position.",
        "Password requirement text is intentionally wrong.",
        "Field labels and button copy differ from Figma.",
        "Input geometry, typography, and panel background are not exact.",
    ],
    "Authentication - Check your email": [
        "Message references SMS instead of email.",
        "Icon is a text glyph rather than the Figma mail icon asset.",
        "Button and resend copy differ from Figma.",
        "Confirmation block is shifted and oversized compared with the original.",
    ],
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9E2EF"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)


def format_run(run, size=None, bold=False, color=None):
    run.font.name = "Calibri"
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    format_run(run, size=22, bold=True, color="0B2545")
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(14)
    r = sub.add_run(subtitle)
    format_run(r, size=11, color="555555")


def add_h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(text)


def add_h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.add_run(text)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_note(doc, label, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, "B8C7DA")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    format_run(r, bold=True, color="1F4D78")
    r = p.add_run(body)
    format_run(r, color="333333")
    doc.add_paragraph()


def setup_doc(title):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run(title).font.size = Pt(9)
    return doc


def add_screen_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    widths = [2500, 2500, 4360]
    headers = ["Screen", "Route", "Purpose"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, "E8EEF5")
        run = cell.paragraphs[0].add_run(text)
        format_run(run, bold=True, color="0B2545")
    for screen, route, purpose in SCREENS:
        row = table.add_row()
        values = [screen, route, purpose]
        for idx, text in enumerate(values):
            cell = row.cells[idx]
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            format_run(run, size=10.5)


def build_candidate_doc():
    doc = setup_doc("Candidate Brief")
    add_title(
        doc,
        "ABC Company Figma Matching Exercise - Candidate Brief",
        "Use the supplied code and original Figma file to identify and fix UI mismatches.",
    )
    add_note(
        doc,
        "Do not use this as an answer key",
        "This brief explains the task and expectations. It intentionally does not list the known bugs.",
    )

    add_h1(doc, "Objective")
    doc.add_paragraph(
        "You are given a Next.js fixture app and the original ABC Company Figma file. Your task is to inspect each route, compare it against the matching Figma screen, and fix the implementation until it visually and behaviorally matches the design."
    )

    add_h1(doc, "Screens to Review")
    add_screen_table(doc)

    add_h1(doc, "What to Fix")
    for item in [
        "Layout geometry: page size, section offsets, grid widths, spacing, and alignment.",
        "Typography: font family, size, weight, line-height, tracking, and text wrapping.",
        "Colors and effects: backgrounds, borders, shadows, opacity, and active states.",
        "Content fidelity: headings, labels, helper text, table values, timers, and button text.",
        "Components: header, avatar, tags, buttons, cards, inputs, editor, modal, sidebars, and tables.",
        "Responsive behavior: screens should remain usable and visually coherent at common desktop and mobile widths.",
    ]:
        add_bullet(doc, item)

    add_h1(doc, "Expected Workflow")
    for item in [
        "Open the route index at http://localhost:3020/.",
        "Open the corresponding Figma frame for each screen.",
        "Create a mismatch list before editing.",
        "Make focused code changes in the fixture app.",
        "Verify each screen in browser after changes.",
        "Submit a short summary of fixes and any remaining uncertainty.",
    ]:
        add_bullet(doc, item)

    add_h1(doc, "Evaluation Criteria")
    for item in [
        "Visual accuracy against the Figma source.",
        "Ability to notice small but material design mismatches.",
        "Clean, maintainable React and CSS changes.",
        "No broad rewrites that hide or bypass the actual mismatch work.",
        "Clear communication of what changed and how it was verified.",
    ]:
        add_bullet(doc, item)

    doc.save(OUT / "InterviewOS_Candidate_Brief.docx")


def build_reviewer_doc():
    doc = setup_doc("Reviewer Guide")
    add_title(
        doc,
        "ABC Company Figma Matching Exercise - Reviewer Guide",
        "Intentional mismatch inventory for scoring candidate submissions.",
    )
    add_note(
        doc,
        "Reviewer use only",
        "This document is the answer key. Do not share it with candidates before or during the exercise.",
    )

    add_h1(doc, "Fixture Scope")
    doc.add_paragraph(
        "The fixture contains all implemented routes from the ABC Company exercise. The implementation intentionally includes visual, content, interaction, and responsive mismatches. Candidates are expected to compare against the original Figma file, not against this guide."
    )
    add_screen_table(doc)

    add_h1(doc, "Screen-by-Screen Mismatch Inventory")
    for screen, route, _ in SCREENS:
        add_h2(doc, f"{screen} ({route})")
        for bug in BUGS[screen]:
            add_bullet(doc, bug)

    add_h1(doc, "Scoring Guidance")
    scoring = doc.add_table(rows=1, cols=3)
    scoring.alignment = WD_TABLE_ALIGNMENT.LEFT
    scoring.autofit = False
    set_table_borders(scoring)
    widths = [1700, 1700, 5960]
    headers = ["Area", "Weight", "Signals"]
    for idx, text in enumerate(headers):
        cell = scoring.rows[0].cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, "E8EEF5")
        r = cell.paragraphs[0].add_run(text)
        format_run(r, bold=True, color="0B2545")
    rows = [
        ("Detection", "35%", "Finds major and subtle mismatches across multiple screen types."),
        ("Implementation", "35%", "Fixes dimensions, content, styling, and layout without brittle hacks."),
        ("Verification", "20%", "Checks routes in browser and explains validation clearly."),
        ("Communication", "10%", "Summarizes changes, tradeoffs, and remaining gaps concisely."),
    ]
    for row_values in rows:
        row = scoring.add_row()
        for idx, text in enumerate(row_values):
            cell = row.cells[idx]
            set_cell_width(cell, widths[idx])
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            format_run(r, size=10.5)

    add_h1(doc, "Reviewer Notes")
    for item in [
        "A strong candidate does not need to catch every seeded issue, but should identify patterns and fix representative issues cleanly.",
        "Prioritize candidates who verify actual rendered output rather than editing from code inspection only.",
        "Watch for changes that make one viewport match while breaking another.",
        "The route index is fixture-only and should not be treated as part of the product Figma design.",
    ]:
        add_bullet(doc, item)

    doc.save(OUT / "InterviewOS_Reviewer_Guide.docx")


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    build_candidate_doc()
    build_reviewer_doc()
    print(OUT / "InterviewOS_Candidate_Brief.docx")
    print(OUT / "InterviewOS_Reviewer_Guide.docx")
